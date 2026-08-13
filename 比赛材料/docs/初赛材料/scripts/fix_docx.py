#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""对 pandoc 产出的 docx 做中文排版处理：字体 / 页面 / 图片居中 / 表格 / 页码"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = sys.argv[1]
# 排版密度档：normal（宽松，阅读优先）/ compact（紧凑，压页数）
DENSITY = sys.argv[2] if len(sys.argv) > 2 else "compact"
doc = Document(PATH)

SONG = "宋体"       # 正文中文
HEI = "黑体"        # 标题中文
LATIN = "Times New Roman"
MONO = "Consolas"

D = {
    "normal": dict(
        body=10.5, line=1.4, after=6.0,
        h=(18, 15, 13, 11.5, 11, 10.5), h_before=12, h_after=6,
        tbl=9.5, tbl_line=1.15, tbl_pad=2.0,
        margin_lr=2.5, margin_t=2.4, margin_b=2.2,
        tall_img=14.0, cap=9.5,
    ),
    # 提交档：在 20 页上限内尽量保住可读性（实测主体 19 页）
    "submit": dict(
        body=10.5, line=1.26, after=3.8,
        h=(16.5, 13.5, 12, 11, 10.5, 10.5), h_before=9, h_after=4,
        tbl=9.0, tbl_line=1.08, tbl_pad=1.3,
        margin_lr=2.2, margin_t=2.1, margin_b=1.9,
        tall_img=13.4, cap=9.0,
    ),
    "compact": dict(
        body=10.0, line=1.18, after=3.0,
        h=(15, 12.5, 11.5, 10.5, 10, 10), h_before=8, h_after=3,
        tbl=8.5, tbl_line=1.02, tbl_pad=0.8,
        margin_lr=2.1, margin_t=2.0, margin_b=1.8,
        tall_img=12.2, cap=8.5,
    ),
}[DENSITY]


def set_fonts(style, latin, east, size=None, bold=None, color=None):
    """同时设置 ascii / hAnsi / eastAsia 三套字体，避免 Word 回退成方框"""
    f = style.font
    f.name = latin
    if size is not None:
        f.size = Pt(size)
    if bold is not None:
        f.bold = bold
    if color is not None:
        f.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:cs"), latin)


names = {s.name for s in doc.styles}

# ── 正文 ──────────────────────────────────────────────
normal = doc.styles["Normal"]
set_fonts(normal, LATIN, SONG, size=D["body"])
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = D["line"]
pf.space_after = Pt(D["after"])
pf.first_line_indent = Cm(0)

# ── 标题 ──────────────────────────────────────────────
_hc = ["8C2318", "1F2937", "1F2937", "374151", "374151", "374151"]
heading_cfg = {f"Heading {i+1}": (D["h"][i], True, _hc[i]) for i in range(6)}
for hname, (sz, bold, color) in heading_cfg.items():
    if hname in names:
        st = doc.styles[hname]
        set_fonts(st, LATIN, HEI, size=sz, bold=bold, color=color)
        st.paragraph_format.space_before = Pt(D["h_before"])
        st.paragraph_format.space_after = Pt(D["h_after"])
        st.paragraph_format.keep_with_next = True

# ── 其他样式 ──────────────────────────────────────────
for nm, (lat, ea, sz) in {
    "Title": (LATIN, HEI, D["h"][0] + 6),
    "Subtitle": (LATIN, HEI, D["h"][2]),
    "Author": (LATIN, SONG, D["body"]),
    "Date": (LATIN, SONG, D["body"]),
    "Quote": (LATIN, SONG, D["body"]),
    "Block Text": (LATIN, SONG, D["body"]),
    "Source Code": (MONO, SONG, D["body"] - 1.5),
    "Verbatim Char": (MONO, SONG, D["body"] - 1),
    "Table Caption": (LATIN, HEI, D["cap"]),
    "Image Caption": (LATIN, SONG, D["cap"]),
    "Caption": (LATIN, SONG, D["cap"]),
    "Compact": (LATIN, SONG, D["body"]),
    "First Paragraph": (LATIN, SONG, D["body"]),
    "Body Text": (LATIN, SONG, D["body"]),
    "Footer": (LATIN, SONG, D["cap"] - 0.5),
    "Header": (LATIN, SONG, D["cap"] - 0.5),
    "List Paragraph": (LATIN, SONG, D["body"]),
}.items():
    if nm in names:
        try:
            set_fonts(doc.styles[nm], lat, ea, size=sz)
        except Exception:
            pass

# 列表样式
for nm in names:
    if nm.startswith(("List", "Bullet", "Compact", "Definition")):
        try:
            set_fonts(doc.styles[nm], LATIN, SONG, size=D["body"])
        except Exception:
            pass

# ── 页面设置：A4 ──────────────────────────────────────
for sec in doc.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(D["margin_lr"])
    sec.right_margin = Cm(D["margin_lr"])
    sec.top_margin = Cm(D["margin_t"])
    sec.bottom_margin = Cm(D["margin_b"])

# ── 图片尺寸：按版心宽度等比缩放 ───────────────────────
TEXT_W = Cm(21.0 - 2 * D["margin_lr"] - 0.2)   # 版心宽度留 2mm 余量
TALL_W = Cm(D["tall_img"])                      # 竖长图收窄，避免独占一页留下大片空白
for shape in doc.inline_shapes:
    aspect = shape.height / shape.width
    target = TALL_W if aspect > 0.7 else TEXT_W
    if shape.width != target:
        ratio = target / shape.width
        shape.height = int(shape.height * ratio)
        shape.width = target

# ── 图片居中 + 图注居中 ───────────────────────────────
img_count = 0
paras = doc.paragraphs
for i, p in enumerate(paras):
    if p._element.findall(".//" + qn("w:drawing")):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True   # 图与图注不跨页拆散
        p.paragraph_format.keep_together = True
        img_count += 1
        # 让图前的标题与引言段一起跟随，避免上一页留大片空白
        for back in range(1, 3):
            if i - back >= 0:
                paras[i - back].paragraph_format.keep_with_next = True
        # 图注：紧随图片段落之后、以「图 N」开头的段落
        if i + 1 < len(paras):
            nxt = paras[i + 1]
            if nxt.style.name in ("Image Caption", "Caption") or nxt.text.strip().startswith("图 "):
                nxt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                nxt.paragraph_format.space_before = Pt(2)
                nxt.paragraph_format.space_after = Pt(12)
                for run in nxt.runs:
                    run.font.size = Pt(D["cap"])
                    run.font.bold = False
                    run.font.color.rgb = RGBColor.from_string("5A5347")
                    rpr = run._element.get_or_add_rPr()
                    rf = rpr.find(qn("w:rFonts"))
                    if rf is None:
                        rf = OxmlElement("w:rFonts")
                        rpr.append(rf)
                    rf.set(qn("w:ascii"), LATIN)
                    rf.set(qn("w:hAnsi"), LATIN)
                    rf.set(qn("w:eastAsia"), HEI)

# ── 表格：加边框 + 表头底色 + 自动宽度 ─────────────────
def set_cell_bg(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFB49A")
        borders.append(el)
    tbl_pr.append(borders)


def text_weight(s):
    """估算文本占宽：中日韩字符按 2 个单位，其余按 1 个"""
    return sum(2 if ord(c) > 0x2E7F else 1 for c in s.strip())


def fit_columns(table, total_twips):
    """按各列内容量按比例分配列宽，避免 pandoc 均分导致内容列被挤窄"""
    ncol = len(table.columns)
    if ncol == 0:
        return
    # 每列取「表头权重」与「正文行平均权重」的较大值，表头不被压断
    weights = []
    for c in range(ncol):
        cells = [r.cells[c].text for r in table.rows]
        head = text_weight(cells[0]) if cells else 1
        body = [text_weight(x) for x in cells[1:]] or [1]
        weights.append(max(head * 0.9, sum(body) / len(body), 4))
    # 压缩极端差距：开平方使宽列不至于吃掉全部空间
    weights = [w ** 0.62 for w in weights]
    lo, hi = 0.07, 0.44          # 单列宽度上下限（占表宽比例）
    frac = [w / sum(weights) for w in weights]
    frac = [min(max(f, lo), hi) for f in frac]
    frac = [f / sum(frac) for f in frac]
    widths = [int(total_twips * f) for f in frac]

    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for i, gc in enumerate(grid.findall(qn("w:gridCol"))):
            if i < len(widths):
                gc.set(qn("w:w"), str(widths[i]))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                tcpr = cell._tc.get_or_add_tcPr()
                for old in tcpr.findall(qn("w:tcW")):
                    tcpr.remove(old)
                tcw = OxmlElement("w:tcW")
                tcw.set(qn("w:w"), str(widths[i]))
                tcw.set(qn("w:type"), "dxa")
                tcpr.append(tcw)


# 版心宽度（twips）：1 cm ≈ 566.9 twips
TBL_TWIPS = int((21.0 - 2 * D["margin_lr"]) * 566.93)

for t in doc.tables:
    set_table_borders(t)
    t.autofit = False
    fit_columns(t, TBL_TWIPS)
    for r_idx, row in enumerate(t.rows):
        for cell in row.cells:
            if r_idx == 0:
                set_cell_bg(cell, "F3EDE0")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(D["tbl_pad"])
                p.paragraph_format.space_after = Pt(D["tbl_pad"])
                p.paragraph_format.line_spacing = D["tbl_line"]
                for run in p.runs:
                    run.font.size = Pt(D["tbl"])
                    rpr = run._element.get_or_add_rPr()
                    rf = rpr.find(qn("w:rFonts"))
                    if rf is None:
                        rf = OxmlElement("w:rFonts")
                        rpr.append(rf)
                    rf.set(qn("w:ascii"), LATIN)
                    rf.set(qn("w:hAnsi"), LATIN)
                    rf.set(qn("w:eastAsia"), SONG)
                    if r_idx == 0:
                        run.font.bold = True

# ── 页脚页码 ──────────────────────────────────────────
def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    for txt, typ in [("begin", "w:fldChar"), (" PAGE ", "w:instrText"),
                     ("end", "w:fldChar")]:
        el = OxmlElement(typ)
        if typ == "w:fldChar":
            el.set(qn("w:fldCharType"), txt)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = txt
        run._r.append(el)
    rpr = run._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), LATIN)
    rf.set(qn("w:eastAsia"), SONG)
    rpr.append(rf)
    run.font.size = Pt(D["cap"] - 0.5)


for sec in doc.sections:
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.text = ""
    add_page_number(fp)

doc.save(PATH)
print(f"✅ 排版完成[{DENSITY}]：{img_count} 张图居中，{len(doc.tables)} 个表格加边框，页脚页码已加")
